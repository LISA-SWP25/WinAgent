import argparse
import getpass
import logging
import os
import random
import subprocess
import threading
import time
import uuid
from datetime import datetime
from pathlib import Path

import requests
import yaml

from actions import apps, net
from client.server_api import send_activity, download_agent_config
from utils.logger import setup_logger

# === Парсим аргументы ===
parser = argparse.ArgumentParser(description="LISA Agent")
parser.add_argument("--debug", action="store_true", help="Run agent in debug mode and clean logs before start")
args = parser.parse_args()

# === Если debug-режим, чистим логи ===
LOG_FILE = Path("agent.log")  # или тот путь, который у тебя реально
if args.debug:
    # Очистка логов
    if LOG_FILE.exists():
        LOG_FILE.unlink()
        print("[*] Debug mode: old logs cleared.")

    # Очистка всех .lock файлов
    for lock_file in Path().glob("*.lock"):
        try:
            lock_file.unlink()
            print(f"[*] Debug mode: removed lock file {lock_file.name}")
        except Exception as e:
            print(f"[!] Failed to remove lock file {lock_file.name}: {e}")

# === Генерация уникального AGENT_ID ===

username = getpass.getuser()
mac_int = uuid.getnode()
mac_str = f"{mac_int:012x}"
AGENT_ID = f"agent_{username}_{mac_str}".lower()

# === Инициализация логгера ===
logger = setup_logger()

# === Singleton: файл-блокировка ===

LOCK_FILE = Path(f"{AGENT_ID}.lock")


def check_singleton():
    if LOCK_FILE.exists():
        logger.error(f"The agent with the ID {AGENT_ID} has already been launched! Completion.")
        sys.exit(1)
    else:
        LOCK_FILE.touch()
        logger.info(f"A blocking file has been created {LOCK_FILE}")


def cleanup_singleton():
    if LOCK_FILE.exists():
        LOCK_FILE.unlink()
        logger.info(f"The lock file {LOCK_FILE} has been deleted")


# === Пути к локальным конфигам (если используются) ===

CONFIG_PATH = Path("config/settings.yaml")
PATHS_PATH = Path("config/paths.yaml")
ROLES_DIR = Path("roles")


def load_yaml(path):
    with open(path, encoding='utf-8') as f:
        return yaml.safe_load(f)


def load_config():
    settings = load_yaml(CONFIG_PATH)
    paths = load_yaml(PATHS_PATH)
    role_name = settings["role"]
    role_file = ROLES_DIR / f"{role_name}.yaml"
    role = load_yaml(role_file)
    return settings, paths, role


# === Выбор случайной активности по весам ===

def weighted_choice(activities):
    weights = [a.get("weight", 1) for a in activities]
    return random.choices(activities, weights=weights, k=1)[0]


# === Основной исполнитель действий ===

import sys


def run_action(action, paths):
    action_type = action.get("action")

    # Задержка перед выполнением, если указана
    delay = action.get("delay")
    if delay:
        logging.info(f"Delay of {delay} seconds before the action {action_type}")
        time.sleep(delay)

    if action_type == "open_app":
        app_name = action.get("app")
        path = paths.get("apps", {}).get(app_name)

        if path == "ms-settings:":
            apps.open_settings()
        elif "dsa.msc" in path:
            apps.open_ad_utilities()
        elif path:
            apps.open_app(path)
        else:
            logging.warning(f"Unknown application: {app_name}")

    elif action_type == "open_browser":
        urls = action.get("urls", ["https://example.com"])
        url = random.choice(urls)
        apps.open_browser(url)

    elif action_type == "run_terminal_command":
        terminal = action.get("terminal")
        command = action.get("command")
        term_path = paths.get("apps", {}).get(terminal)
        if term_path:
            apps.run_terminal_command(term_path, command)
        else:
            logging.warning(f"Unknown terminal: {terminal}")

    elif action_type == "sleep":
        seconds = action.get("seconds", 60)
        logging.info(f"Pause (sleep) for {seconds} seconds")
        time.sleep(seconds)

    elif action_type == "simulate_activity":
        net.simulate_network_activity()

    elif action_type == "terminate":
        logging.info("The agent is being terminated according to the script.")
        sys.exit(0)

    elif action_type == "os_settings":
        apps.open_settings()

    elif action_type == "ad_utilities":
        apps.open_ad_utilities()

    elif action_type == "create_file":
        path = os.path.expandvars(action.get("path", ""))
        content = action.get("content", "")
        editor = action.get("editor")

        if not path:
            logger.warning("No path specified for create_file")
            return

        try:
            if editor == "word" and path.endswith(".docx"):
                from docx import Document
                doc = Document()
                doc.add_paragraph(content)
                doc.save(path)
                logger.info(f"DOCX file created: {path}")
            else:
                with open(path, "w", encoding="utf-8") as f:
                    f.write(content)
                logger.info(f"File created: {path}")

            if editor == "notepad":
                subprocess.Popen(["notepad.exe", path])
            elif editor == "vscode":
                raw_vscode_path = paths.get("apps", {}).get("vscode", "code")
                vscode_path = os.path.expandvars(raw_vscode_path)
                subprocess.Popen([vscode_path, path])
            elif editor == "word":
                raw_word_path = paths.get("apps", {}).get("word", "WINWORD.EXE")
                word_path = os.path.expandvars(raw_word_path)
                subprocess.Popen([word_path, path])

        except Exception as e:
            logger.error(f"Failed to create file: {e}")

    elif action_type == "read_file":
        path = os.path.expandvars(action.get("path", ""))
        editor = action.get("editor")

        if not path:
            logger.warning("No path specified for read_file")
            return

        try:
            if os.path.exists(path):
                if editor == "notepad":
                    subprocess.Popen(["notepad.exe", path])
                elif editor == "vscode":
                    raw_vscode_path = paths.get("apps", {}).get("vscode", "code")
                    vscode_path = os.path.expandvars(raw_vscode_path)
                    subprocess.Popen([vscode_path, path])
                elif editor == "word":
                    raw_word_path = paths.get("apps", {}).get("word", "WINWORD.EXE")
                    word_path = os.path.expandvars(raw_word_path)
                    subprocess.Popen([word_path, path])

                logger.info(f"File opened for reading: {path}")
            else:
                logger.warning(f"File not found for reading: {path}")

        except Exception as e:
            logger.error(f"Failed to read file: {e}")

    elif action_type == "update_file":
        path = os.path.expandvars(action.get("path", ""))
        content = action.get("content", "")
        editor = action.get("editor")

        if not path:
            logger.warning("No path specified for update_file")
            return

        try:
            if os.path.exists(path):
                if editor == "word" and path.endswith(".docx"):
                    from docx import Document
                    doc = Document(path)
                    if doc.paragraphs:
                        # Добавляем к последнему параграфу
                        doc.paragraphs[-1].add_run(content)
                    else:
                        # Если пусто — создаём первый параграф
                        doc.add_paragraph(content)
                    doc.save(path)
                    logger.info(f"DOCX file updated (appended): {path}")
                else:
                    with open(path, "a", encoding="utf-8") as f:
                        f.write(content)
                    logger.info(f"File updated (appended): {path}")
            else:
                logger.warning(f"File not found for update: {path}")

        except Exception as e:
            logger.error(f"Failed to update file: {e}")


    elif action_type == "delete_file":
        path = os.path.expandvars(action.get("path", ""))

        if not path:
            logger.warning("No path specified for delete_file")
            return

        try:
            if os.path.exists(path):
                os.remove(path)
                logger.info(f"File deleted: {path}")
            else:
                logger.warning(f"File not found for deletion: {path}")

        except Exception as e:
            logger.error(f"Failed to delete file: {e}")


    elif action_type == "conditional_exit":
        condition = action.get("condition")
        # Пример: проверка нерабочего времени
        if condition == "not_work_time" and not is_work_time({}):  # передай settings если нужно
            logging.info("The condition is met: the agent is shutting down.")
            sys.exit(0)
        else:
            logging.info("The completion condition is not met — the agent continues to work.")

    else:
        logging.error(f"Unknown action: {action_type}")


# === Проверка рабочего времени ===

def is_work_time(settings):
    now = datetime.now()
    weekday = now.isoweekday()
    current_time = now.time()
    work_days = settings["work_days"]
    start = datetime.strptime(settings["work_start"], "%H:%M").time()
    end = datetime.strptime(settings["work_end"], "%H:%M").time()
    return weekday in work_days and start <= current_time <= end


def user_exists(username):
    result = subprocess.run(
        ["net", "user", username],
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True
    )
    return "The user name could not be found" not in result.stdout


def send_heartbeat():
    while True:
        payload = {
            "agent_id": AGENT_ID,
            "username": username,
            "timestamp": datetime.utcnow().isoformat(),
            "role": "Admin",
            "system_info": {
                "platform": os.name
            },
            "status": "active"
        }

        # Используй свой URL и API ключ, если есть
        url = "http://localhost:8000/api/agents/heartbeat"
        headers = {"Authorization": "Bearer sk-agent-heartbeat-key-2024"}

        try:
            response = requests.post(url, json=payload, headers=headers, timeout=10)
            response.raise_for_status()
            logger.info(f"Heartbeat sent: {response.json()}")
        except Exception as e:
            logger.error(f"Failed to send heartbeat: {e}")

        # Спим 24 часа = 86400 сек
        time.sleep(86400)


# === Основной цикл ===


def main():
    check_singleton()
    logger.info("Launching the LISA agent with the configuration by ID")

    try:
        config = download_agent_config(AGENT_ID)
        if not config:
            logger.error("Agent config was not received. Completion.")
            return

        interval = config.get("custom_config", {}).get("interval", 10)
        randomize = config.get("custom_config", {}).get("randomize", True)
        tasks = config.get("behavior_template", {}).get("tasks", [])

        if not tasks:
            logger.error("The task template is empty. Completion.")
            return

        logger.info(f"Agent: {AGENT_ID}, Interval: {interval}, Tasks: {tasks}")
        settings, paths, _ = load_config()

        # === Heartbeat поток ===
        t_heartbeat = threading.Thread(target=send_heartbeat, daemon=True)
        t_heartbeat.start()

        # === Запускаем repeatable задачи в фоне ===
        for task in tasks:
            if task.get("repeatable"):
                t = threading.Thread(
                    target=run_repeatable, args=(task, paths), daemon=True
                )
                t.start()

        if randomize:
            while True:
                if not user_exists(username):
                    logger.info(f"User '{username}' no longer exists. The agent is stopping.")
                    break

                task = weighted_choice([t for t in tasks if not t.get("repeatable")])
                run_action(task, paths=paths)

                send_activity(AGENT_ID, task, {"status": "ok"})
                logger.info(f"Waiting {interval} seconds")
                time.sleep(interval)
        else:
            while True:
                for task in tasks:
                    if task.get("repeatable"):
                        continue  # Skip, уже крутится в фоне

                    if not user_exists(username):
                        logger.info(f"User '{username}' no longer exists. The agent is stopping.")
                        break

                    run_action(task, paths=paths)

                    send_activity(AGENT_ID, task, {"status": "ok"})
                    logger.info(f"Waiting {interval} seconds")
                    time.sleep(interval)

    finally:
        cleanup_singleton()


def run_repeatable(task, paths):
    while True:
        run_action(task, paths)
        delay = task.get("delay", 5)
        logger.info(f"Repeatable task '{task['action']}' sleeping {delay} sec")
        time.sleep(delay)


# === Запуск ===

if __name__ == "__main__":
    main()
